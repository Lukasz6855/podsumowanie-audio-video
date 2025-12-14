from moviepy.editor import VideoFileClip  # Biblioteka do edycji wideo
import os  # Operacje na systemie plików
from pathlib import Path  # Obsługa ścieżek plików
from fpdf import FPDF  # Biblioteka do tworzenia plików PDF
from docx import Document  # Biblioteka do tworzenia plików DOCX
from docx.shared import Pt, Inches  # Jednostki miary dla DOCX
import io  # Operacje wejścia/wyjścia na bajtach
import yt_dlp  # Biblioteka do pobierania filmów z YouTube
import tempfile  # Tworzenie plików tymczasowych

def pobierz_wideo_z_youtube(url):
    """
    Funkcja pobierająca wideo z YouTube i zapisująca je jako plik MP4.
    
    Args:
        url (str): Link do filmu na YouTube
        
    Returns:
        tuple: (sciezka_do_pliku, nazwa_pliku) - ścieżka do pobranego pliku i jego tytuł
    """
    # Utworzenie katalogu tymczasowego dla pobranych filmów
    katalog_temp = tempfile.mkdtemp()
    
    # Konfiguracja yt-dlp
    opcje_ydl = {
        'format': 'bestvideo[ext=mp4]+bestaudio[ext=m4a]/best[ext=mp4]/best',  # Najlepsza jakość MP4
        'outtmpl': os.path.join(katalog_temp, '%(title)s.%(ext)s'),  # Szablon nazwy pliku
        'quiet': False,  # Wyświetlanie postępu
        'no_warnings': False,  # Wyświetlanie ostrzeżeń
    }
    
    # Pobranie wideo
    with yt_dlp.YoutubeDL(opcje_ydl) as ydl:
        # Pobranie informacji o wideo
        info = ydl.extract_info(url, download=True)
        
        # Pobranie tytułu i nazwy pliku
        tytul = info.get('title', 'video')
        
        # Wyczyszczenie tytułu z niedozwolonych znaków
        tytul_bezpieczny = "".join(c for c in tytul if c.isalnum() or c in (' ', '-', '_')).strip()
        
        # Znalezienie pobranego pliku
        sciezka_pliku = ydl.prepare_filename(info)
    
    # Zwrócenie ścieżki do pliku i tytułu
    return sciezka_pliku, tytul_bezpieczny

def wyodrebnij_audio_z_wideo(sciezka_wideo):
    """
    Funkcja wyodrębniająca audio z pliku wideo i zapisująca jako MP3.
    
    Args:
        sciezka_wideo (str): Ścieżka do pliku wideo
        
    Returns:
        str: Ścieżka do wygenerowanego pliku audio MP3
    """
    # Utworzenie ścieżki dla pliku wyjściowego audio
    sciezka_audio = sciezka_wideo.rsplit('.', 1)[0] + '_audio.mp3'  # Zamiana rozszerzenia na .mp3
    
    # Wczytanie pliku wideo
    klip_wideo = VideoFileClip(sciezka_wideo)  # Utworzenie obiektu klip wideo
    
    # Wyodrębnienie ścieżki audio
    klip_audio = klip_wideo.audio  # Pobranie komponentu audio z wideo
    
    # Zapisanie audio jako plik MP3
    klip_audio.write_audiofile(sciezka_audio, codec='mp3')  # Zapis audio w formacie MP3
    
    # Zamknięcie klipów aby zwolnić zasoby
    klip_audio.close()  # Zamknięcie klip audio
    klip_wideo.close()  # Zamknięcie klip wideo
    
    # Zwrócenie ścieżki do zapisanego pliku audio
    return sciezka_audio


def oblicz_koszt_transkrypcji(dlugosc_minuty):
    """
    Funkcja obliczająca koszt transkrypcji przy użyciu Whisper-1.
    
    Args:
        dlugosc_minuty (float): Długość audio w minutach
        
    Returns:
        float: Koszt transkrypcji w dolarach
    """
    # Whisper-1 kosztuje $0.006 za minutę
    koszt_za_minute = 0.006  # Koszt w USD za minutę transkrypcji
    
    # Obliczenie całkowitego kosztu
    return dlugosc_minuty * koszt_za_minute


def oblicz_koszt_gpt(tokeny_wejsciowe, tokeny_wyjsciowe):
    """
    Funkcja obliczająca szacunkowy koszt użycia GPT-4o.
    
    Args:
        tokeny_wejsciowe (int): Liczba tokenów w zapytaniu (input)
        tokeny_wyjsciowe (int): Liczba tokenów w odpowiedzi (output)
        
    Returns:
        float: Szacunkowy koszt w dolarach
    """
    # GPT-4o koszty (przykładowe stawki - należy zaktualizować według aktualnego cennika)
    koszt_input_za_1k = 0.0025  # Koszt za 1000 tokenów wejściowych w USD
    koszt_output_za_1k = 0.01  # Koszt za 1000 tokenów wyjściowych w USD
    
    # Obliczenie kosztu dla tokenów wejściowych
    koszt_input = (tokeny_wejsciowe / 1000) * koszt_input_za_1k
    
    # Obliczenie kosztu dla tokenów wyjściowych
    koszt_output = (tokeny_wyjsciowe / 1000) * koszt_output_za_1k
    
    # Zwrócenie łącznego kosztu
    return koszt_input + koszt_output


def pobierz_rozmiar_pliku_mb(sciezka_pliku):
    """
    Funkcja pobierająca rozmiar pliku w megabajtach.
    
    Args:
        sciezka_pliku (str): Ścieżka do pliku
        
    Returns:
        float: Rozmiar pliku w megabajtach
    """
    # Pobranie rozmiaru pliku w bajtach
    rozmiar_bajtow = os.path.getsize(sciezka_pliku)
    
    # Konwersja bajtów na megabajty (1 MB = 1024 * 1024 bajtów)
    rozmiar_mb = rozmiar_bajtow / (1024 * 1024)
    
    # Zwrócenie rozmiaru w MB
    return rozmiar_mb


def pobierz_dlugosc_audio(sciezka_pliku):
    """
    Funkcja pobierająca długość pliku audio/wideo w minutach.
    
    Args:
        sciezka_pliku (str): Ścieżka do pliku audio lub wideo
        
    Returns:
        float: Długość pliku w minutach
    """
    from moviepy.editor import VideoFileClip, AudioFileClip  # Import bibliotek do obsługi multimediów
    
    rozszerzenie = Path(sciezka_pliku).suffix.lower()  # Pobranie rozszerzenia pliku
    
    try:  # Próba odczytu długości
        if rozszerzenie in ['.mp4', '.avi', '.mov']:  # Jeśli plik wideo
            klip = VideoFileClip(sciezka_pliku)  # Wczytanie wideo
            dlugosc_sekund = klip.duration  # Pobranie długości w sekundach
            klip.close()  # Zamknięcie klip
        else:  # Jeśli plik audio
            klip = AudioFileClip(sciezka_pliku)  # Wczytanie audio
            dlugosc_sekund = klip.duration  # Pobranie długości w sekundach
            klip.close()  # Zamknięcie klip
        
        # Konwersja sekund na minuty
        dlugosc_minut = dlugosc_sekund / 60
        return dlugosc_minut  # Zwrócenie długości w minutach
    except Exception as e:  # Obsługa błędów
        return 0  # Zwrócenie 0 w przypadku błędu


def formatuj_czas_na_min_sec(dlugosc_minuty):
    """
    Funkcja formatująca długość z formatu minut (float) na format "X min Y sec".
    
    Args:
        dlugosc_minuty (float): Długość w minutach (np. 5.75 dla 5 minut i 45 sekund)
        
    Returns:
        str: Sformatowany czas w formacie "X min Y sec" (np. "5 min 45 sec")
    """
    # Sprawdzenie czy długość jest poprawna
    if dlugosc_minuty is None or dlugosc_minuty <= 0:  # Jeśli długość niepoprawna
        return "0 min 0 sec"  # Zwrócenie domyślnej wartości
    
    # Obliczenie pełnych minut
    pelne_minuty = int(dlugosc_minuty)  # Rzutowanie na int aby uzyskać pełne minuty
    
    # Obliczenie sekund w kolejnej rozpoczętej minucie
    sekundy = int((dlugosc_minuty - pelne_minuty) * 60)  # Obliczenie pozostałych sekund
    
    # Zwrócenie sformatowanego czasu
    return f"{pelne_minuty} min {sekundy} sec"  # Format: X min Y sec


def zlicz_slowa(tekst):
    """
    Funkcja zliczająca liczbę słów w tekście.
    
    Args:
        tekst (str): Tekst do analizy
        
    Returns:
        int: Liczba słów w tekście
    """
    # Podzielenie tekstu na słowa i zliczenie
    slowa = tekst.split()  # Podział na słowa
    return len(slowa)  # Zwrócenie liczby słów


def szacuj_tokeny_z_slow(liczba_slow):
    """
    Funkcja szacująca liczbę tokenów na podstawie liczby słów.
    
    Args:
        liczba_slow (int): Liczba słów
        
    Returns:
        int: Szacowana liczba tokenów (1 słowo ≈ 1.3 tokena dla języka polskiego)
    """
    # Dla języka polskiego przyjmujemy że 1 słowo ≈ 1.3 tokena
    return int(liczba_slow * 1.3)


def generuj_plik_txt(nazwa_pliku, transkrypcja, podsumowanie):
    """
    Funkcja generująca zawartość pliku TXT z transkrypcją i podsumowaniem.
    
    Args:
        nazwa_pliku (str): Nazwa oryginalnego pliku
        transkrypcja (str): Tekst transkrypcji
        podsumowanie (str): Tekst podsumowania
        
    Returns:
        str: Zawartość pliku TXT
    """
    # Tworzenie zawartości pliku tekstowego
    zawartosc = f"""TRANSKRYPCJA I PODSUMOWANIE
{'=' * 50}

Plik źródłowy: {nazwa_pliku}

{'=' * 50}
TRANSKRYPCJA
{'=' * 50}

{transkrypcja}

{'=' * 50}
PODSUMOWANIE
{'=' * 50}

{podsumowanie}
"""
    # Zwrócenie zawartości
    return zawartosc


def generuj_plik_pdf(nazwa_pliku, transkrypcja, podsumowanie):
    """
    Funkcja generująca plik PDF z transkrypcją i podsumowaniem.
    
    Args:
        nazwa_pliku (str): Nazwa oryginalnego pliku
        transkrypcja (str): Tekst transkrypcji
        podsumowanie (str): Tekst podsumowania
        
    Returns:
        bytes: Zawartość pliku PDF w postaci bajtów
    """
    # Funkcja pomocnicza do transliteracji polskich znaków
    def transliteruj_tekst(tekst):  # Zamiana polskich znaków na odpowiedniki ASCII
        mapa_znakow = {  # Słownik mapowania polskich znaków
            'ą': 'a', 'ć': 'c', 'ę': 'e', 'ł': 'l', 'ń': 'n', 
            'ó': 'o', 'ś': 's', 'ź': 'z', 'ż': 'z',
            'Ą': 'A', 'Ć': 'C', 'Ę': 'E', 'Ł': 'L', 'Ń': 'N',
            'Ó': 'O', 'Ś': 'S', 'Ź': 'Z', 'Ż': 'Z'
        }
        for pol, ascii in mapa_znakow.items():  # Iteracja po mapowaniu
            tekst = tekst.replace(pol, ascii)  # Zamiana znaku
        return tekst  # Zwrócenie przetworzonego tekstu
    
    # Utworzenie obiektu PDF
    pdf = FPDF()  # Inicjalizacja obiektu PDF
    pdf.add_page()  # Dodanie nowej strony
    
    # Ustawienie czcionki (używamy domyślnej Arial)
    pdf.set_font("Arial", 'B', 16)  # Czcionka pogrubiona, rozmiar 16
    
    # Dodanie tytułu
    pdf.cell(0, 10, transliteruj_tekst('TRANSKRYPCJA I PODSUMOWANIE'), 0, 1, 'C')  # Tytuł wyśrodkowany
    pdf.ln(5)  # Odstęp pionowy
    
    # Dodanie nazwy pliku źródłowego
    pdf.set_font("Arial", '', 12)  # Czcionka normalna, rozmiar 12
    pdf.cell(0, 10, transliteruj_tekst(f'Plik zrodlowy: {nazwa_pliku}'), 0, 1)  # Nazwa pliku
    pdf.ln(5)  # Odstęp pionowy
    
    # Sekcja transkrypcji
    pdf.set_font("Arial", 'B', 14)  # Czcionka pogrubiona, rozmiar 14
    pdf.cell(0, 10, 'TRANSKRYPCJA', 0, 1)  # Nagłówek transkrypcji
    pdf.ln(2)  # Odstęp pionowy
    
    pdf.set_font("Arial", '', 11)  # Czcionka normalna, rozmiar 11
    # Podzielenie transkrypcji na linie i dodanie do PDF
    transkrypcja_ascii = transliteruj_tekst(transkrypcja)  # Transliteracja całej transkrypcji
    for linia in transkrypcja_ascii.split('\n'):  # Iteracja po liniach
        if linia.strip():  # Sprawdzenie czy linia nie jest pusta
            pdf.multi_cell(0, 6, linia)  # Komórka z automatycznym zawijaniem tekstu
        else:  # Jeśli linia pusta
            pdf.ln(3)  # Dodanie odstępu
    
    pdf.ln(5)  # Odstęp pionowy
    
    # Sekcja podsumowania
    pdf.set_font("Arial", 'B', 14)  # Czcionka pogrubiona, rozmiar 14
    pdf.cell(0, 10, 'PODSUMOWANIE', 0, 1)  # Nagłówek podsumowania
    pdf.ln(2)  # Odstęp pionowy
    
    pdf.set_font("Arial", '', 11)  # Czcionka normalna, rozmiar 11
    # Podzielenie podsumowania na linie i dodanie do PDF
    podsumowanie_ascii = transliteruj_tekst(podsumowanie)  # Transliteracja całego podsumowania
    for linia in podsumowanie_ascii.split('\n'):  # Iteracja po liniach
        if linia.strip():  # Sprawdzenie czy linia nie jest pusta
            pdf.multi_cell(0, 6, linia)  # Komórka z automatycznym zawijaniem tekstu
        else:  # Jeśli linia pusta
            pdf.ln(3)  # Dodanie odstępu
    
    # Zwrócenie PDF jako bajty
    pdf_output = pdf.output(dest='S')  # Generowanie PDF jako string
    if isinstance(pdf_output, str):  # Jeśli wynik jest stringiem
        return pdf_output.encode('latin-1', errors='replace')  # Konwersja do bajtów z obsługą błędów
    else:  # Jeśli wynik jest już bajtami
        return pdf_output  # Zwrócenie bezpośrednio


def generuj_plik_docx(nazwa_pliku, transkrypcja, podsumowanie):
    """
    Funkcja generująca plik DOCX z transkrypcją i podsumowaniem.
    
    Args:
        nazwa_pliku (str): Nazwa oryginalnego pliku
        transkrypcja (str): Tekst transkrypcji
        podsumowanie (str): Tekst podsumowania
        
    Returns:
        bytes: Zawartość pliku DOCX w postaci bajtów
    """
    # Utworzenie nowego dokumentu Word
    dokument = Document()  # Inicjalizacja dokumentu
    
    # Dodanie tytułu
    naglowek = dokument.add_heading('TRANSKRYPCJA I PODSUMOWANIE', 0)  # Nagłówek poziomu 0
    naglowek.alignment = 1  # Wyśrodkowanie (1 = center)
    
    # Dodanie nazwy pliku źródłowego
    dokument.add_paragraph(f'Plik źródłowy: {nazwa_pliku}')  # Paragraf z nazwą pliku
    
    # Dodanie sekcji transkrypcji
    dokument.add_heading('TRANSKRYPCJA', level=1)  # Nagłówek poziomu 1
    paragraf_transkrypcji = dokument.add_paragraph(transkrypcja)  # Paragraf z transkrypcją
    paragraf_transkrypcji.style.font.size = Pt(11)  # Ustawienie rozmiaru czcionki
    
    # Dodanie odstępu
    dokument.add_paragraph()  # Pusty paragraf jako odstęp
    
    # Dodanie sekcji podsumowania
    dokument.add_heading('PODSUMOWANIE', level=1)  # Nagłówek poziomu 1
    paragraf_podsumowania = dokument.add_paragraph(podsumowanie)  # Paragraf z podsumowaniem
    paragraf_podsumowania.style.font.size = Pt(11)  # Ustawienie rozmiaru czcionki
    
    # Zapisanie dokumentu do bufora bajtów
    bufor = io.BytesIO()  # Utworzenie bufora w pamięci
    dokument.save(bufor)  # Zapisanie dokumentu do bufora
    bufor.seek(0)  # Przesunięcie wskaźnika na początek bufora
    
    # Zwrócenie zawartości bufora jako bajty
    return bufor.getvalue()


def generuj_audio_z_tekstu(klient, tekst, sciezka_wyjsciowa, glos="onyx"):
    """
    Funkcja generująca plik audio z tekstu przy użyciu OpenAI TTS.
    
    Args:
        klient: Klient OpenAI API
        tekst (str): Tekst do zamiany na mowę
        sciezka_wyjsciowa (str): Ścieżka gdzie zapisać wygenerowany plik MP3
        glos (str): Głos do użycia (alloy, echo, fable, onyx, nova, shimmer)
        
    Returns:
        str: Ścieżka do wygenerowanego pliku audio
    """
    # Wywołanie API OpenAI TTS
    odpowiedz = klient.audio.speech.create(
        model="tts-1",  # Model TTS (tańszy i szybszy)
        voice=glos,  # Wybrany głos
        input=tekst  # Tekst do zamiany na mowę
    )
    
    # Zapisanie odpowiedzi audio do pliku
    odpowiedz.stream_to_file(sciezka_wyjsciowa)  # Zapis do pliku MP3
    
    # Zwrócenie ścieżki do pliku
    return sciezka_wyjsciowa


def oblicz_koszt_tts(liczba_znakow):
    """
    Funkcja obliczająca koszt generowania mowy z tekstu przy użyciu TTS-1.
    
    Args:
        liczba_znakow (int): Liczba znaków w tekście
        
    Returns:
        float: Koszt w dolarach
    """
    # TTS-1 kosztuje $0.015 za 1000 znaków
    koszt_za_1k_znakow = 0.015  # Koszt w USD za 1000 znaków
    
    # Obliczenie całkowitego kosztu
    return (liczba_znakow / 1000) * koszt_za_1k_znakow
